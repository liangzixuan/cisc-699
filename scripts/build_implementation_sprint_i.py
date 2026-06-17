#!/usr/bin/env python3
"""Build the W5 implementation sprint I check-in DOCX from Markdown."""

from __future__ import annotations

import re
import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PKG = ROOT / "docs" / "05-implementation-sprint-i"
DEFAULT_SRC = PKG / "implementation-sprint-i-check-in.md"
DEFAULT_OUT = PKG / "Implementation-Sprint-I-Check-in.docx"

COLORS = {
    "navy": "17365D",
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "gray": "F2F4F7",
    "border": "A6A6A6",
}


def set_document_style(doc: Document, footer_label: str = "SafeExec W5 Check-in") -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Title", 22, COLORS["navy"], 0, 10),
        ("Heading 1", 16, COLORS["blue"], 16, 8),
        ("Heading 2", 13, COLORS["blue"], 12, 6),
        ("Heading 3", 12, COLORS["dark_blue"], 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(footer_label)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_inline_markdown(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Menlo"
            run.font.size = Pt(9)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    raw_rows: list[str] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw_rows.append(lines[i].strip())
        i += 1

    rows: list[list[str]] = []
    for raw in raw_rows:
        cells = [cell.strip() for cell in raw.strip("|").split("|")]
        if all(set(cell) <= {"-", ":", " "} for cell in cells):
            continue
        rows.append(cells)
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return

    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True

    for row_index, row in enumerate(rows):
        for col_index in range(cols):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, COLORS["gray"])
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline_markdown(paragraph, row[col_index] if col_index < len(row) else "")
            for run in paragraph.runs:
                run.font.size = Pt(8.4)
                if row_index == 0:
                    run.bold = True
    doc.add_paragraph()


def add_code_block(doc: Document, code: list[str]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.right_indent = Inches(0.25)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run("\n".join(code))
    run.font.name = "Menlo"
    run.font.size = Pt(8.5)


def build(
    source: Path = DEFAULT_SRC,
    output: Path = DEFAULT_OUT,
    footer_label: str = "SafeExec W5 Check-in",
) -> None:
    doc = Document()
    set_document_style(doc, footer_label)
    lines = source.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_buffer: list[str] = []

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_buffer)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        if stripped.startswith("# "):
            paragraph = doc.add_paragraph(style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_markdown(paragraph, stripped[2:])
            i += 1
            continue
        if stripped.startswith("## "):
            paragraph = doc.add_paragraph(style="Heading 1")
            add_inline_markdown(paragraph, stripped[3:])
            i += 1
            continue
        if stripped.startswith("### "):
            paragraph = doc.add_paragraph(style="Heading 2")
            add_inline_markdown(paragraph, stripped[4:])
            i += 1
            continue
        if stripped.startswith("> "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.right_indent = Inches(0.25)
            add_inline_markdown(paragraph, stripped[2:])
            for run in paragraph.runs:
                run.italic = True
            i += 1
            continue
        if stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(paragraph, stripped[2:])
            i += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.32)
            paragraph.paragraph_format.first_line_indent = Inches(-0.24)
            paragraph.add_run(f"{numbered.group(1)}.  ")
            add_inline_markdown(paragraph, numbered.group(2))
            i += 1
            continue

        paragraph = doc.add_paragraph()
        add_inline_markdown(paragraph, stripped)
        i += 1

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(output)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Build W5 DOCX artifacts.")
    parser.add_argument("--source", type=Path, default=DEFAULT_SRC)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUT)
    parser.add_argument("--footer-label", default="SafeExec W5 Check-in")
    args = parser.parse_args()
    build(args.source, args.output, args.footer_label)
