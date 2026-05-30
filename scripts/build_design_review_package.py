#!/usr/bin/env python3
"""Build the W4 design review package DOCX and diagrams from Markdown."""

from __future__ import annotations

import re
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
PKG = ROOT / "docs" / "04-design-review-package"
SRC = PKG / "design-review-package.md"
OUT = PKG / "Design-Review-Package.docx"
DIAGRAM_DIR = PKG / "diagrams"


COLORS = {
    "navy": "#17365D",
    "blue": "#2E74B5",
    "light_blue": "#D9EAF7",
    "green": "#E3F4E4",
    "purple": "#EAE4F8",
    "red": "#F8E5E5",
    "gray": "#F2F4F7",
    "line": "#4A5568",
}


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def draw_box(draw: ImageDraw.ImageDraw, xy, text: str, fill: str, outline: str = "#4A5568", width: int = 2) -> None:
    draw.rounded_rectangle(xy, radius=14, fill=fill, outline=outline, width=width)
    x1, y1, x2, y2 = xy
    wrapped = []
    for part in text.split("\n"):
        wrapped.extend(textwrap.wrap(part, width=max(12, int((x2 - x1) / 12))) or [""])
    line_h = 18
    total_h = line_h * len(wrapped)
    y = y1 + (y2 - y1 - total_h) / 2
    for line in wrapped:
        bbox = draw.textbbox((0, 0), line, font=font(15, bold=True))
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, fill="#111827", font=font(15, bold=True))
        y += line_h


def arrow(draw: ImageDraw.ImageDraw, start, end, label: str | None = None) -> None:
    draw.line([start, end], fill=COLORS["line"], width=3)
    sx, sy = start
    ex, ey = end
    dx = ex - sx
    dy = ey - sy
    if abs(dx) >= abs(dy):
        pts = [(ex, ey), (ex - 12 if dx > 0 else ex + 12, ey - 7), (ex - 12 if dx > 0 else ex + 12, ey + 7)]
    else:
        pts = [(ex, ey), (ex - 7, ey - 12 if dy > 0 else ey + 12), (ex + 7, ey - 12 if dy > 0 else ey + 12)]
    draw.polygon(pts, fill=COLORS["line"])
    if label:
        mx = (sx + ex) / 2
        my = (sy + ey) / 2
        draw.text((mx - 45, my - 22), label, fill="#374151", font=font(13))


def create_diagrams() -> None:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)

    img = Image.new("RGB", (1500, 850), "white")
    d = ImageDraw.Draw(img)
    d.text((50, 30), "SafeExec Component Architecture", fill=COLORS["navy"], font=font(30, True))

    draw_box(d, (80, 140, 300, 230), "Client / Test Harness", COLORS["light_blue"])
    draw_box(d, (390, 125, 610, 245), "FastAPI\nPOST /execute\nPydantic models", COLORS["green"])
    draw_box(d, (700, 125, 930, 245), "Execution\nOrchestrator\nworkspace + timeout", COLORS["green"])
    draw_box(d, (1040, 80, 1320, 185), "DockerExecutor\nhardened Docker controls", COLORS["purple"])
    draw_box(d, (1040, 235, 1320, 340), "GVisorExecutor\nrunsc runtime", COLORS["purple"])
    draw_box(d, (700, 395, 930, 500), "Ephemeral\nWorkspace\nper request", COLORS["gray"])
    draw_box(d, (1040, 430, 1320, 535), "Result Artifacts\nJSONL / CSV\nhost metadata", COLORS["gray"])
    draw_box(d, (390, 610, 610, 720), "Functional +\nAdversarial\nCorpora", COLORS["red"])
    draw_box(d, (700, 610, 930, 720), "Benchmark\nRunner\n>=30 samples", COLORS["red"])

    arrow(d, (300, 185), (390, 185), "JSON")
    arrow(d, (610, 185), (700, 185), "request")
    arrow(d, (930, 165), (1040, 135), "selects")
    arrow(d, (930, 215), (1040, 285), "selects")
    arrow(d, (815, 245), (815, 395), "mount")
    arrow(d, (930, 460), (1040, 480), "writes")
    arrow(d, (500, 610), (500, 245), "API calls")
    arrow(d, (815, 610), (815, 500), "measures")
    arrow(d, (700, 665), (610, 665), None)
    img.save(DIAGRAM_DIR / "component-architecture.png")

    img = Image.new("RGB", (1500, 850), "white")
    d = ImageDraw.Draw(img)
    d.text((50, 30), "Evaluation Flow and Evidence Collection", fill=COLORS["navy"], font=font(30, True))

    draw_box(d, (80, 120, 330, 230), "Functional Suite\nHumanEval / MBPP\nstudent-authored", COLORS["light_blue"])
    draw_box(d, (80, 310, 330, 420), "Adversarial Suite\n>=40 programs\n>=6 categories", COLORS["red"])
    draw_box(d, (80, 500, 330, 610), "Benchmark Suite\ncold / warm /\nsteady-state", COLORS["gray"])
    draw_box(d, (470, 220, 720, 370), "Shared API\nPOST /execute", COLORS["green"])
    draw_box(d, (850, 120, 1110, 230), "Docker Results\npass rate\ncontainment\nlatency", COLORS["purple"])
    draw_box(d, (850, 430, 1110, 540), "gVisor Results\npass rate\ncontainment\nlatency", COLORS["purple"])
    draw_box(d, (1230, 260, 1450, 400), "Report Evidence\nCSV/JSON\ntables + CIs\nlimitations", COLORS["green"])

    for y in [175, 365, 555]:
        arrow(d, (330, y), (470, 295), "requests")
    arrow(d, (720, 270), (850, 175), "backend A")
    arrow(d, (720, 320), (850, 485), "backend B")
    arrow(d, (1110, 175), (1230, 305), None)
    arrow(d, (1110, 485), (1230, 355), None)
    img.save(DIAGRAM_DIR / "evaluation-flow.png")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_document_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Title", 22, "17365D", 0, 10),
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_inline_markdown(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()].replace("*", ""))
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Menlo"
            run.font.size = Pt(9)
        elif token.startswith("["):
            label = token[1 : token.index("]")]
            href = token[token.index("(") + 1 : -1]
            run = paragraph.add_run(f"{label} ({href})")
            run.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:].replace("*", ""))


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    table_lines: list[str] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        table_lines.append(lines[i].strip())
        i += 1
    rows = []
    for raw in table_lines:
        cells = [cell.strip() for cell in raw.strip("|").split("|")]
        if all(set(cell) <= {"-", ":", " "} for cell in cells):
            continue
        rows.append(cells)
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, COLORS["gray"])
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline_markdown(paragraph, row[c_idx] if c_idx < len(row) else "")
            for run in paragraph.runs:
                run.font.size = Pt(8.2)
                if r_idx == 0:
                    run.bold = True
    doc.add_paragraph()


def add_image(doc: Document, alt: str, rel_path: str) -> None:
    image_path = (SRC.parent / rel_path).resolve()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(5.7))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = caption.add_run(alt)
    cap_run.italic = True
    cap_run.font.size = Pt(9)


def build() -> None:
    create_diagrams()
    doc = Document()
    set_document_style(doc)
    lines = SRC.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_buffer: list[str] = []

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                para = doc.add_paragraph()
                run = para.add_run("\n".join(code_buffer))
                run.font.name = "Menlo"
                run.font.size = Pt(8.5)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        image_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            add_image(doc, image_match.group(1), image_match.group(2))
            i += 1
            continue
        if stripped.startswith("# "):
            p = doc.add_paragraph(style="Title")
            add_inline_markdown(p, stripped[2:])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        if stripped.startswith("## "):
            if stripped == "## References":
                doc.add_section(WD_SECTION_START.NEW_PAGE)
            p = doc.add_paragraph(style="Heading 1")
            add_inline_markdown(p, stripped[3:])
            i += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline_markdown(p, stripped[4:])
            i += 1
            continue
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.25)
            add_inline_markdown(p, stripped[2:])
            for run in p.runs:
                run.italic = True
            i += 1
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(p, stripped[2:])
            i += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.32)
            p.paragraph_format.first_line_indent = Inches(-0.24)
            p.add_run(f"{numbered.group(1)}.  ")
            add_inline_markdown(p, numbered.group(2))
            i += 1
            continue
        p = doc.add_paragraph()
        add_inline_markdown(p, stripped)
        i += 1

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
