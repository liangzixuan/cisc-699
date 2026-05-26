#!/usr/bin/env python3
"""Build the W3 literature/requirements brief DOCX from Markdown."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "docs" / "03-lit-req-brief" / "literature-and-requirements-brief.md"
OUT = ROOT / "docs" / "03-lit-req-brief" / "Literature-and-Requirements-Brief.docx"


def set_document_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.05

    for name, size, before, after in [
        ("Title", 22, 0, 8),
        ("Heading 1", 15, 12, 5),
        ("Heading 2", 12, 9, 4),
        ("Heading 3", 10.5, 7, 3),
    ]:
        style = doc.styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_inline_markdown(paragraph, text: str) -> None:
    text = text.replace("\\`", "`")
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()].replace("*", ""))
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Menlo"
            run.font.size = Pt(8.5)
        elif token.startswith("["):
            label = token[1 : token.index("]")]
            href = token[token.index("(") + 1 : -1]
            run = paragraph.add_run(f"{label} ({href})")
            run.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:].replace("*", ""))


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    table_lines: list[str] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        table_lines.append(lines[i].strip())
        i += 1
    rows = []
    for raw in table_lines:
        cells = [cell.strip() for cell in raw.strip("|").split("|")]
        if all(set(cell) <= {"-", ":", " "} for cell in cells):
            continue
        rows.append(cells)
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            value = row[c_idx] if c_idx < len(row) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline_markdown(paragraph, value)
            for run in paragraph.runs:
                run.font.size = Pt(8)
                if r_idx == 0:
                    run.bold = True
    doc.add_paragraph()


def add_image(doc: Document, alt: str, rel_path: str) -> None:
    image_path = (SRC.parent / rel_path).resolve()
    if not image_path.exists():
        paragraph = doc.add_paragraph()
        paragraph.add_run(f"[Missing image: {rel_path}]").italic = True
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.7))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.add_run(alt).italic = True


def build() -> None:
    doc = Document()
    set_document_style(doc)

    lines = SRC.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_buffer: list[str] = []

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                para = doc.add_paragraph(style="No Spacing")
                run = para.add_run("\n".join(code_buffer))
                run.font.name = "Menlo"
                run.font.size = Pt(8)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        image_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            add_image(doc, image_match.group(1), image_match.group(2))
            i += 1
            continue

        if stripped.startswith("# "):
            title = stripped[2:].strip()
            paragraph = doc.add_paragraph(style="Title")
            add_inline_markdown(paragraph, title)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if stripped.startswith("## "):
            if stripped == "## References":
                doc.add_section(WD_SECTION_START.NEW_PAGE)
            paragraph = doc.add_paragraph(style="Heading 1")
            add_inline_markdown(paragraph, stripped[3:].strip())
            i += 1
            continue

        if stripped.startswith("### "):
            paragraph = doc.add_paragraph(style="Heading 2")
            add_inline_markdown(paragraph, stripped[4:].strip())
            i += 1
            continue

        if stripped.startswith("> "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.right_indent = Inches(0.25)
            add_inline_markdown(paragraph, stripped[2:].strip())
            for run in paragraph.runs:
                run.italic = True
            i += 1
            continue

        if stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(paragraph, stripped[2:].strip())
            i += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered:
            paragraph = doc.add_paragraph(style="List Number")
            add_inline_markdown(paragraph, numbered.group(2))
            i += 1
            continue

        paragraph = doc.add_paragraph()
        add_inline_markdown(paragraph, stripped)
        i += 1

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
